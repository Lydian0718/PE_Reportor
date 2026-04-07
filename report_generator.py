from docxtpl import DocxTemplate
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import datetime
import io
import sys

def resource_path(relative_path):
    """获取资源的绝对路径，兼容开发环境和 PyInstaller 打包环境"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

LOGO_PATH = "assets/Logo_big.png"

def get_doc_template(template_path=None):
    if template_path is None:
        template_path = resource_path("template_with_header.docx")
    return DocxTemplate(template_path)

def setup_document(doc):
    # 渲染页眉
    context = {'header_company': "嘉兴市南湖股权投资基金有限公司"}
    doc.render(context)
    
    document = doc.docx
    # 清除原有正文
    for paragraph in list(document.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)
    for table in list(document.tables):
        t = table._element
        t.getparent().remove(t)
    return document

def add_paragraph(document, text, style_func=None, indent=False, align=None):
    p = document.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = p.add_run(text)
    if style_func:
        style_func(run)
    return p

def set_big_title_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run.font.size = Pt(36)
    run.bold = True

def set_title_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run.font.size = Pt(14)
    run.bold = True

def set_heading1_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(12)
    run.bold = True

def set_heading2_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run.font.size = Pt(10.5)
    run.bold = True

def set_body_font(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run.font.size = Pt(10.5)
    run.bold = False

def generate_brief_report(data):
    doc = get_doc_template()
    document = setup_document(doc)
    
    # 标题
    add_paragraph(document, f"{data['company_short_name']}项目研判报告", set_title_font, align='center')
    document.add_paragraph()
    
    # 一、公司概况
    add_paragraph(document, "一、公司概况", set_heading1_font)
    add_paragraph(document, data["company_overview"], set_body_font, indent=True, align='justify')
    document.add_paragraph()
    
    # 二、融资情况
    add_paragraph(document, "二、融资情况", set_heading1_font)
    add_paragraph(document, data["financing"], set_body_font, indent=True, align='justify')
    document.add_paragraph()
    
    # 三、项目亮点与风险
    add_paragraph(document, "三、项目亮点与风险", set_heading1_font)
    add_paragraph(document, "（一）项目亮点", set_heading2_font)
    for i, item in enumerate(data["highlights"], 1):
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.3)
        run_num = p.add_run(f"（{i}）")
        set_body_font(run_num); run_num.bold = True
        run1 = p.add_run(item["title"])
        set_body_font(run1); run1.bold = True
        run2 = p.add_run(item["detail"])
        set_body_font(run2)
    document.add_paragraph()
    
    add_paragraph(document, "（二）项目风险", set_heading2_font)
    for i, item in enumerate(data["risks"], 1):
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.3)
        run_num = p.add_run(f"（{i}）")
        set_body_font(run_num); run_num.bold = True
        run1 = p.add_run(item["title"])
        set_body_font(run1); run1.bold = True
        run2 = p.add_run(item["detail"])
        set_body_font(run2)
    document.add_paragraph()
    
    # 四、投资建议
    add_paragraph(document, "四、投资建议", set_heading1_font)
    add_paragraph(document, data["investment_opinion"], set_body_font, indent=True, align='justify')
    document.add_paragraph()
    
    # 决策建议
    p = document.add_paragraph()
    run = p.add_run(f"决策建议：{data['decision'].upper()}")
    set_body_font(run)
    
    file_stream = io.BytesIO()
    doc.docx.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_detailed_report(data):
    doc = get_doc_template()
    document = setup_document(doc)
    
    # 封面
    for _ in range(4): document.add_paragraph()
    actual_logo = resource_path(LOGO_PATH)
    if os.path.exists(actual_logo):
        for _ in range(2): document.add_paragraph()
        p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(); run.add_picture(actual_logo, width=Inches(3), height=Inches(3))
    for _ in range(2): document.add_paragraph()
    add_paragraph(document, f"{data['company_short_name']}项目报告", set_big_title_font, align='center')
    for _ in range(2): document.add_paragraph()
    year_month = datetime.datetime.now().strftime("%Y年%m月")
    add_paragraph(document, year_month, set_title_font, align='center')
    document.add_page_break()
    
    # 正文
    add_paragraph(document, f"{data['company_short_name']}项目分析报告", set_title_font, align='center')
    document.add_paragraph()

    sections = [
        ("一、公司概况", "company_overview"),
        ("二、行业概况", "industry_overview"),
        ("三、业务概况", "business_overview"),
        ("四、财务概况", "financial_overview"),
        ("五、融资情况", "financing"),
        ("六、上市前景分析", "listing_analysis")
    ]
    for title, key in sections:
        add_paragraph(document, title, set_heading1_font)
        add_paragraph(document, data[key], set_body_font, indent=True, align='justify')
        document.add_paragraph()
    
    add_paragraph(document, "七、项目亮点与风险", set_heading1_font)
    for sub_title, key in [("（一）项目亮点", "highlights"), ("（二）项目风险", "risks")]:
        add_paragraph(document, sub_title, set_heading2_font)
        for i, item in enumerate(data[key], 1):
            p = document.add_paragraph(); p.paragraph_format.first_line_indent = Inches(0.3)
            run_num = p.add_run(f"（{i}）"); set_body_font(run_num); run_num.bold = True
            run1 = p.add_run(item["title"]); set_body_font(run1); run1.bold = True
            run2 = p.add_run(item["detail"]); set_body_font(run2)
        document.add_paragraph()
        
    add_paragraph(document, "八、投资建议", set_heading1_font)
    add_paragraph(document, data["investment_opinion"], set_body_font, indent=True, align='justify')
    document.add_paragraph()
    
    add_paragraph(document, "九、尽调关注要点", set_heading1_font)
    for i, item in enumerate(data["key_concerns"], 1):
        p = document.add_paragraph(); p.paragraph_format.first_line_indent = Inches(0.3)
        run_num = p.add_run(f"（{i}）"); set_body_font(run_num); run_num.bold = True
        run1 = p.add_run(item["title"]); set_body_font(run1); run1.bold = True
        run2 = p.add_run(item["detail"]); set_body_font(run2)
    document.add_paragraph()
    
    p = document.add_paragraph()
    run = p.add_run(f"决策建议：{data['decision'].upper()}"); set_body_font(run)
    
    file_stream = io.BytesIO()
    doc.docx.save(file_stream)
    file_stream.seek(0)
    return file_stream